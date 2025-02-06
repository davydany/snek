import click
import io
import markdown
import openai
import os
import pdfkit             # Or use PyPDF or other libraries
import streamlit as st
import tempfile

from docx import Document  # For .docx export


# Convert notes to desired formats
def generate_markdown_export():
    output = []
    if 'notes' in st.session_state and len(st.session_state["notes"]) > 0:
        for note_obj in st.session_state["notes"]:
            output.append(f"### Section {note_obj['chunk_idx']+1}\n")
            if note_obj["summary"]:
                output.append(f"**AI Summary**:\n{note_obj['summary']}\n\n")
            if note_obj["note"]:
                output.append(f"**My Notes**:\n{note_obj['note']}\n\n")
        return "\n".join(output)
    else:
        return ""

def generate_docx():
    doc = Document()
    if 'notes' in st.session_state and len(st.session_state["notes"]) > 0:
        for note_obj in st.session_state["notes"]:
            doc.add_heading(f"Section {note_obj['chunk_idx']+1}", level=2)
            if note_obj["summary"]:
                doc.add_paragraph("AI Summary:")
                doc.add_paragraph(note_obj["summary"])
            if note_obj["note"]:
                doc.add_paragraph("My Notes:")
                doc.add_paragraph(note_obj["note"])
    return doc
    


st.set_page_config(
    page_title="Research Aide",
    page_icon="🧑‍🔬️",
    layout="wide",
    initial_sidebar_state="auto",
)

# Setup the App Title
st.title("🧑‍🔬️ Research Aide")

upload_col, settings_col = st.columns([0.7, 0.3])
with upload_col:

    st.subheader("Upload Notes")

    # File Uploader
    uploaded_file = st.file_uploader(
        "Upload DOCX or PDF", 
        type=["docx", "pdf"],
        accept_multiple_files=False
    )
    if uploaded_file is not None:
        st.session_state["uploaded_file"] = uploaded_file

    st.divider()

with settings_col:

    with st.expander("Settings"):
        openai_api_key = st.text_input("OpenAI API Key", value=os.getenv("OPENAI_API_KEY"))
    
    markdown_col, docx_col, pdf_col, html_col = st.columns(4)

    st.subheader("Export Notes")
    with markdown_col:
        export_md = generate_markdown_export()
        st.download_button(
            label="⬇️ Markdown",
            data=export_md,
            file_name="notes_export.md",
            mime="text/markdown"
        )

    with docx_col:
        
        docx_file = generate_docx()
        docx_buffer = io.BytesIO()
        docx_file.save(docx_buffer)
        st.download_button(
            label="⬇️ DOCX",
            data=docx_buffer.getvalue(),
            file_name="notes_export.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    with pdf_col:

        pdf_data = pdfkit.from_string(export_md, False)  
        st.download_button(
            label="⬇️ PDF",
            data=pdf_data,
            file_name="notes_export.pdf",
            mime="application/pdf"
        )

    with html_col:
        html_data = markdown.markdown(export_md)
        st.download_button(
            label="⬇️ HTML",
            data=html_data,
            file_name="notes_export.html",
            mime="text/html"
        )

# Initialize session_state for notes
if "notes" not in st.session_state:
    st.session_state["notes"] = []

if "md_content" not in st.session_state:
    st.session_state.md_content = ""

if uploaded_file is not None:

    with st.spinner("🔍 Wait for it..."):

        # Save the uploaded file to disk
        click.secho("Saving file to disk...", fg='green')
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_file.write(uploaded_file.getvalue())
            temp_file_name = temp_file.name

        # Convert to Markdown with Docling
        click.secho("Converting to Markdown...", fg='green')
        from docling.document_converter import DocumentConverter
        converter = DocumentConverter()
        result = converter.convert(source=temp_file_name)
        md_content = result.document.export_to_markdown()
        if st.session_state.get("md_content", "") != md_content:

            st.session_state.md_content = md_content


        # Split content by paragraphs, tables, images, etc.
        # (Pseudo-code: docling may provide chunking or parse method)
        chunks = md_content.split("\n\n")  # simplistic paragraph split example
        count = 0

        for idx, chunk in enumerate(chunks):

            click.secho("Iterating thru chunk #" + str(idx), fg='green')

            # check and see if the chunk has enough text,
            # because if it's not, it's not worth the effort
            # to show the user and summarize it, or take 
            # notes on it
            if len(chunk) < 100:
                continue

            count += 1
            with st.container(border=True):
                
                # Two-column layout
                col_left, col_right = st.columns(2)

                # Display original doc (Markdown) on the left
                with col_left:
                    st.markdown(chunk)

                # On the right, show text boxes for notes + AI summary button
                with col_right:
                    st.markdown("## Summaries & Notes")
                    st.write(f"**Section {count}**")

                    # Summarize Button
                    if st.button(f"Summarize Section {count}"):
                        # Replace with your own OpenAI logic
                        openai_client = openai.OpenAI(openai_api_key)
                        response = openai_client.Completion.create(
                            model="text-davinci-003",
                            prompt=f"Summarize this:\n\n{chunk}",
                            max_tokens=60
                        )
                        summary_text = response.choices[0].text.strip()
                        st.session_state["notes"].append({"chunk_idx": idx, 
                                                        "summary": summary_text, 
                                                        "note": ""})

                    # Display existing summary
                    summary = ""
                    for note_obj in st.session_state["notes"]:
                        if note_obj["chunk_idx"] == count:
                            summary = note_obj["summary"]

                    st.text_area(f"AI Summary for Section {count}", value=summary, key=f"summary_{idx}")

                    # User notes
                    user_note_key = f"note_{idx}"
                    user_note = st.text_area(f"Your Notes for Section {count}", key=user_note_key)

                    # Keep notes in session_state
                    # (Could update on every keystroke or only on a save button.)
                    if st.button(f"Save Notes for Section {count}"):
                        # Update or create an entry
                        updated = False
                        for note_obj in st.session_state["notes"]:
                            if note_obj["chunk_idx"] == count:
                                note_obj["note"] = st.session_state[user_note_key]
                                updated = True
                        if not updated:
                            st.session_state["notes"].append({
                                "chunk_idx": count,
                                "summary": "",
                                "note": st.session_state[user_note_key]
                            })
                        st.success("Notes saved.")
                st.divider()

